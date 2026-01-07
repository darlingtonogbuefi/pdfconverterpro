# main.py

import os
import tempfile
from pathlib import Path
from typing import List

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse

# PDF / OCR
from pdf2image import convert_from_path
from PIL import Image
import pytesseract
import fitz  # PyMuPDF

# Word / Excel / PDF
from pdf2docx import Converter as PDFToWordConverter
from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch

# -------------------------------
# CONFIG (Windows)
# -------------------------------
POPPLER_PATH = r"C:\poppler\Library\bin"
TESSERACT_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH

# -------------------------------
# FASTAPI SETUP
# -------------------------------
app = FastAPI(title="Universal File Converter API", version="2.1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# -------------------------------
# HELPERS
# -------------------------------
def encode_file_to_base64(file_path: str) -> str:
    import base64
    with open(file_path, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")


def is_scanned_pdf(pdf_path: str, min_chars: int = 20) -> bool:
    doc = fitz.open(pdf_path)
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()
    return len(text.strip()) < min_chars


def render_pdf_images(pdf_path: str):
    """
    Force REAL Poppler rendering (no Pillow fallback)
    """
    return convert_from_path(
        pdf_path,
        poppler_path=POPPLER_PATH,
        dpi=300,
        fmt="png",
        use_pdftocairo=True
    )


# -------------------------------
# HEALTH
# -------------------------------
@app.get("/health")
async def health():
    return {"status": "ok"}


# -------------------------------
# PDF → WORD (AUTO OCR)
# -------------------------------
@app.post("/api/convert/pdf-to-word")
async def pdf_to_word(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(400, "File must be a PDF")

    try:
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = os.path.join(tmp, "input.pdf")
            docx_path = os.path.join(tmp, "output.docx")

            with open(pdf_path, "wb") as f:
                f.write(await file.read())

            scanned = is_scanned_pdf(pdf_path)

            if not scanned:
                cv = PDFToWordConverter(pdf_path)
                cv.convert(docx_path)
                cv.close()
            else:
                doc = Document()
                images = render_pdf_images(pdf_path)
                for img in images:
                    doc.add_paragraph(pytesseract.image_to_string(img))
                doc.save(docx_path)

            return JSONResponse({
                "success": True,
                "filename": Path(file.filename).stem + (".docx" if not scanned else "_ocr.docx"),
                "file": encode_file_to_base64(docx_path)
            })

    except Exception as e:
        raise HTTPException(500, str(e))


# -------------------------------
# PDF → EXCEL (DIGITAL + SCANNED)
# -------------------------------
import camelot  # pip install camelot-py[cv]

@app.post("/api/convert/pdf-to-excel")
async def pdf_to_excel(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(400, "File must be a PDF")

    try:
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = os.path.join(tmp, "input.pdf")
            xlsx_path = os.path.join(tmp, "output.xlsx")

            # Save uploaded PDF
            with open(pdf_path, "wb") as f:
                f.write(await file.read())

            scanned = is_scanned_pdf(pdf_path)

            if not scanned:
                # DIGITAL PDF: Use Camelot to extract tables
                tables = camelot.read_pdf(pdf_path, pages="all", flavor="lattice")
                if not tables:
                    tables = camelot.read_pdf(pdf_path, pages="all", flavor="stream")

                if not tables:
                    raise HTTPException(400, "No tables found in PDF")

                # Create Excel workbook
                from openpyxl import Workbook
                wb = Workbook()
                ws = wb.active
                ws.title = "Sheet1"

                row_offset = 1
                for table in tables:
                    for r_idx, row in enumerate(table.df.values):
                        for c_idx, cell in enumerate(row):
                            ws.cell(row=row_offset + r_idx, column=c_idx + 1, value=cell)
                    row_offset += len(table.df) + 2  # leave 2 empty rows between tables

                wb.save(xlsx_path)

            else:
                # SCANNED PDF: Use OCR
                from openpyxl import Workbook
                wb = Workbook()
                ws = wb.active
                row = 1

                images = render_pdf_images(pdf_path)
                for img in images:
                    data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
                    current_row = {}
                    last_top = None

                    for i, text in enumerate(data['text']):
                        if not text.strip():
                            continue

                        top = data['top'][i]
                        left = data['left'][i]

                        # Start new row if vertical position changed significantly
                        if last_top is None or abs(top - last_top) > 10:
                            # write previous row
                            if current_row:
                                for col, val in sorted(current_row.items()):
                                    ws.cell(row=row, column=col + 1, value=val)
                                row += 1
                            current_row = {}
                            last_top = top

                        col_idx = left // 100  # crude column based on horizontal position
                        current_row[col_idx] = text

                    # write last row
                    if current_row:
                        for col, val in sorted(current_row.items()):
                            ws.cell(row=row, column=col + 1, value=val)
                        row += 1

                wb.save(xlsx_path)

            return JSONResponse({
                "success": True,
                "filename": Path(file.filename).stem + (".xlsx" if not scanned else "_ocr.xlsx"),
                "file": encode_file_to_base64(xlsx_path)
            })

    except Exception as e:
        raise HTTPException(500, str(e))

# -------------------------------
# PDF → IMAGES (REAL RENDERING)
# -------------------------------
@app.post("/api/convert/pdf-to-images")
async def pdf_to_images(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(400, "File must be a PDF")

    try:
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = os.path.join(tmp, "input.pdf")
            with open(pdf_path, "wb") as f:
                f.write(await file.read())

            images = render_pdf_images(pdf_path)
            encoded_images = []

            for i, img in enumerate(images):
                img = img.convert("RGB")  # safe for PNG/JPG
                img_path = os.path.join(tmp, f"page_{i+1}.png")
                img.save(img_path, "PNG")
                encoded_images.append(encode_file_to_base64(img_path))

            return JSONResponse({
                "success": True,
                "images": encoded_images
            })

    except Exception as e:
        raise HTTPException(500, str(e))


# -------------------------------
# WORD → PDF
# -------------------------------
@app.post("/api/convert/word-to-pdf")
async def word_to_pdf(file: UploadFile = File(...)):
    if not file.filename.lower().endswith((".doc", ".docx")):
        raise HTTPException(400, "Invalid Word file")

    try:
        with tempfile.TemporaryDirectory() as tmp:
            docx_path = os.path.join(tmp, "input.docx")
            pdf_path = os.path.join(tmp, "output.pdf")

            with open(docx_path, "wb") as f:
                f.write(await file.read())

            doc = Document(docx_path)
            c = canvas.Canvas(pdf_path, pagesize=letter)
            y = letter[1] - inch

            for p in doc.paragraphs:
                if p.text.strip():
                    c.drawString(inch, y, p.text)
                    y -= 14
                    if y < inch:
                        c.showPage()
                        y = letter[1] - inch

            c.save()

            return JSONResponse({
                "success": True,
                "filename": Path(file.filename).stem + ".pdf",
                "file": encode_file_to_base64(pdf_path)
            })

    except Exception as e:
        raise HTTPException(500, str(e))


# -------------------------------
# IMAGES → PDF
# -------------------------------
@app.post("/api/convert/images-to-pdf")
async def images_to_pdf(files: List[UploadFile] = File(...)):
    try:
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = os.path.join(tmp, "output.pdf")
            c = canvas.Canvas(pdf_path, pagesize=letter)

            for file in files:
                img_path = os.path.join(tmp, file.filename)
                with open(img_path, "wb") as f:
                    f.write(await file.read())

                img = Image.open(img_path).convert("RGB")
                width = letter[0] - 2 * inch
                height = width * (img.height / img.width)

                c.drawInlineImage(img, inch, inch, width, height)
                c.showPage()

            c.save()

            return JSONResponse({
                "success": True,
                "filename": "images.pdf",
                "file": encode_file_to_base64(pdf_path)
            })

    except Exception as e:
        raise HTTPException(500, str(e))
